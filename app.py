from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, make_response
from docx import Document
from io import BytesIO
import docx.shared
from docx.shared import RGBColor, Pt
from werkzeug.utils import secure_filename
from PIL import Image
import tempfile
import os
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from weasyprint import HTML


app = Flask(__name__)

def appliquer_style(doc, theme):
    """Applique un style au document en fonction du thème choisi"""
    if theme == "moderne":
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
    elif theme == "couleur":
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        doc.styles['Normal'].font.color.rgb = RGBColor(0, 102, 204)
    else:  # classique (par défaut)
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(12)

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/form')
def form():
    return render_template('form.html')

@app.route('/generate', methods=['POST'])
def generate_cv():
    # Vérification du consentement
    if not request.form.get("consentement"):
        return "Vous devez accepter la politique de traitement des données.", 400

    # Récupération des données
    theme = request.form.get("theme", "classique")
    nom = request.form.get("nom", "").strip()
    age = request.form.get("age", "").strip()
    titre = request.form.get("titre", "").strip()
    ville = request.form.get("ville", "").strip()
    email = request.form.get("email", "").strip()
    telephone = request.form.get("telephone", "").strip()
    profil = request.form.get("profil", "").strip()
    experiences = request.form.get("experiences", "").strip()
    competences = request.form.get("competences", "").strip()
    langues = request.form.get("langues", "").strip()
    formations = request.form.get("formations", "").strip()
    interets = request.form.get("interets", "").strip()


    # Validation des champs obligatoires
    if not nom or not age:
        return "Les champs 'Nom' et 'Âge' sont obligatoires.", 400

    # Sauvegarde dans la base SQLite
    conn = sqlite3.connect('cv_data.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS cvs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            age TEXT,
            titre TEXT,
            ville TEXT,
            email TEXT,
            telephone TEXT,
            profil TEXT,
            experiences TEXT,
            competences TEXT,
            langues TEXT,
            formations TEXT,
            interets TEXT
        )
    ''')
    c.execute('''
        INSERT INTO cvs (nom, age, titre, ville, email, telephone, profil, experiences, competences, langues, formations, interets)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (nom, age, titre, ville, email, telephone, profil, experiences, competences, langues, formations, interets))
    conn.commit()
    conn.close()

    # Récupération de la photo
    photo = request.files.get("photo")
    photo_path = None
    if photo:
        try:
            # Sauvegarder temporairement la photo
            temp_dir = tempfile.gettempdir()
            photo_filename = secure_filename(photo.filename)
            photo_path = os.path.join(temp_dir, photo_filename)
            photo.save(photo_path)

            # Convertir l'image en format compatible (JPEG/PNG)
            with Image.open(photo_path) as img:
                if img.format not in ["JPEG", "PNG"]:
                    photo_path = os.path.join(temp_dir, f"{os.path.splitext(photo_filename)[0]}.png")
                    img.convert("RGB").save(photo_path, "PNG")
        except Exception as e:
            return f"Erreur lors du traitement de la photo : {str(e)}", 500

    # Création du document
    doc = Document()
    appliquer_style(doc, theme)

    # Ajouter la photo au document
    if photo_path:
        try:
            doc.add_picture(photo_path, width=docx.shared.Inches(1.5))
        except Exception as e:
            return f"Erreur lors de l'ajout de la photo au document : {str(e)}", 500

    # Supprimer la photo temporaire
    if photo_path:
        try:
            os.remove(photo_path)
        except Exception as e:
            print(f"Erreur lors de la suppression de la photo temporaire : {str(e)}")

    # En-tête du CV
    doc.add_heading(f"{nom} - {age} ans", level=0)
    if titre:
        doc.add_paragraph(titre)
    
    # Coordonnées
    coordonnees = []
    if ville: coordonnees.append(f"📍 {ville}")
    if email: coordonnees.append(f"📧 {email}")
    if telephone: coordonnees.append(f"📞 {telephone}")
    
    if coordonnees:
        doc.add_paragraph("\n".join(coordonnees))

    # Sections du CV
    if profil:
        doc.add_heading("Profil", level=1)
        doc.add_paragraph(profil)

    if experiences:
        doc.add_heading("Expériences professionnelles", level=1)
        for exp in filter(None, (e.strip() for e in experiences.split('\n'))):
            doc.add_paragraph(exp, style='List Bullet')

    if competences:
        doc.add_heading("Compétences", level=1)
        for comp in filter(None, (c.strip() for c in competences.split(','))):
            doc.add_paragraph(f"• {comp}", style='List Bullet')

    if langues:
        doc.add_heading("Langues", level=1)
        for lang in filter(None, (l.strip() for l in langues.split(','))):
            doc.add_paragraph(f"• {lang}", style='List Bullet')

    if formations:
        doc.add_heading("Formation", level=1)
        for form in filter(None, (f.strip() for f in formations.split('\n'))):
            doc.add_paragraph(f"• {form}", style='List Bullet')

    if interets:
        doc.add_heading("Centres d'intérêt", level=1)
        for interest in filter(None, (i.strip() for i in interets.split(','))):
            doc.add_paragraph(f"• {interest}", style='List Bullet')

    format = request.form.get("format", "docx")  # Default format is docx

    if format == "pdf":
        # Generate PDF
        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=letter)
        pdf.setFont("Helvetica", 12)

        # Header with photo
        if photo_path:
            try:
                pdf.drawImage(photo_path, 50, 700, width=100, height=100)  # Add photo at the top-left
            except Exception as e:
                print(f"Erreur lors de l'ajout de la photo au PDF : {e}")

        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(160, 750, f"{nom} - {age} ans")
        if titre:
            pdf.setFont("Helvetica", 14)
            pdf.drawString(160, 730, titre)

        # Contact Information
        y = 710
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(160, y, "Coordonnées :")
        y -= 20
        pdf.setFont("Helvetica", 12)
        if ville:
            pdf.drawString(160, y, f"📍 {ville}")
            y -= 20
        if email:
            pdf.drawString(160, y, f"📧 {email}")
            y -= 20
        if telephone:
            pdf.drawString(160, y, f"📞 {telephone}")
            y -= 20

        # Sections with styled layout
        def add_section(title, content, y):
            if content:
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawString(50, y, title)
                y -= 20
                pdf.setFont("Helvetica", 12)
                for line in content.split('\n'):
                    pdf.drawString(50, y, line.strip())
                    y -= 20
            return y

        y = add_section("Profil", profil, y - 40)
        y = add_section("Expériences professionnelles", experiences, y - 40)
        y = add_section("Compétences", competences.replace(',', '\n'), y - 40)
        y = add_section("Langues", langues.replace(',', '\n'), y - 40)
        y = add_section("Formation", formations, y - 40)
        y = add_section("Centres d'intérêt", interets.replace(',', '\n'), y - 40)

        pdf.save()
        buffer.seek(0)

        # Remove temporary photo
        if photo_path:
            try:
                os.remove(photo_path)
            except Exception as e:
                print(f"Erreur lors de la suppression de la photo temporaire : {e}")

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"CV_{nom.replace(' ', '_')}.pdf",
            mimetype='application/pdf'
        )

    # Génération du fichier
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"CV_{nom.replace(' ', '_')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/cvs')
def liste_cvs():
    conn = sqlite3.connect('cv_data.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS cvs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            age TEXT,
            titre TEXT,
            ville TEXT,
            email TEXT,
            telephone TEXT,
            profil TEXT,
            experiences TEXT,
            competences TEXT,
            langues TEXT,
            formations TEXT,
            interets TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    c.execute("SELECT * FROM cvs ORDER BY created_at DESC")
    rows = c.fetchall()
    conn.close()
    return render_template("liste_cvs.html", cvs=rows)

@app.route('/delete/<int:id>', methods=['POST'])
def delete_cv(id):
    conn = sqlite3.connect("cv_data.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM cvs WHERE id = ?", (id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/save_motivation', methods=['POST'])
def save_motivation():
    data = request.json
    nom = data.get("nom", "").strip()
    adresse = data.get("adresse", "").strip()
    ville = data.get("ville", "").strip()
    email = data.get("email", "").strip()
    tel = data.get("tel", "").strip()
    date_ville = data.get("date_ville", "").strip()
    sujet = data.get("sujet", "").strip()
    competences = data.get("competences", "").strip()
    experience = data.get("experience", "").strip()
    qualites = data.get("qualites", "").strip()

    # Vérifie si une lettre identique existe déjà (évite les doublons sur la même génération)
    conn = sqlite3.connect('cv_data.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS lettres (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            adresse TEXT,
            ville TEXT,
            email TEXT,
            tel TEXT,
            date_ville TEXT,
            sujet TEXT,
            competences TEXT,
            experience TEXT,
            qualites TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    c.execute('''
        SELECT id FROM lettres WHERE
            nom=? AND adresse=? AND ville=? AND email=? AND tel=? AND date_ville=? AND sujet=? AND competences=? AND experience=? AND qualites=?
        ORDER BY created_at DESC LIMIT 1
    ''', (nom, adresse, ville, email, tel, date_ville, sujet, competences, experience, qualites))
    exists = c.fetchone()
    if not exists:
        c.execute('''
            INSERT INTO lettres (nom, adresse, ville, email, tel, date_ville, sujet, competences, experience, qualites)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (nom, adresse, ville, email, tel, date_ville, sujet, competences, experience, qualites))
        conn.commit()
    conn.close()
    return jsonify({"status": "ok"})

@app.route('/admin')
def admin():
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS cvs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            age TEXT,
            titre TEXT,
            ville TEXT,
            email TEXT,
            telephone TEXT,
            profil TEXT,
            experiences TEXT,
            competences TEXT,
            langues TEXT,
            formations TEXT,
            interets TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute("SELECT * FROM cvs ORDER BY created_at DESC")
    rows = cursor.fetchall()
    # Lettres de motivation
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS lettres (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            adresse TEXT,
            ville TEXT,
            email TEXT,
            tel TEXT,
            date_ville TEXT,
            sujet TEXT,
            competences TEXT,
            experience TEXT,
            qualites TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute("SELECT * FROM lettres ORDER BY created_at DESC")
    lettres = cursor.fetchall()
    conn.close()
    return render_template("admin.html", rows=rows, lettres=lettres)

@app.route('/edit/<int:id>', methods=['GET', 'POST'])
def edit_cv(id):
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    if request.method == 'POST':
        # Update the CV entry
        nom = request.form.get("nom", "").strip()
        age = request.form.get("age", "").strip()
        titre = request.form.get("titre", "").strip()
        ville = request.form.get("ville", "").strip()
        email = request.form.get("email", "").strip()
        telephone = request.form.get("telephone", "").strip()
        profil = request.form.get("profil", "").strip()
        experiences = request.form.get("experiences", "").strip()
        competences = request.form.get("competences", "").strip()
        langues = request.form.get("langues", "").strip()
        formations = request.form.get("formations", "").strip()
        interets = request.form.get("interets", "").strip()

        cursor.execute('''
            UPDATE cvs
            SET nom = ?, age = ?, titre = ?, ville = ?, email = ?, telephone = ?, profil = ?, experiences = ?, competences = ?, langues = ?, formations = ?, interets = ?
            WHERE id = ?
        ''', (nom, age, titre, ville, email, telephone, profil, experiences, competences, langues, formations, interets, id))
        conn.commit()
        conn.close()
        return "CV mis à jour avec succès.", 200

    # Fetch the CV entry for editing
    cursor.execute("SELECT * FROM cvs WHERE id = ?", (id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return "CV introuvable.", 404
    return render_template("edit.html", cv=row)

def generate_professional_cv(doc, data):
    """Generate a professional CV using a predefined template."""
    # Add header with name and title
    header = doc.add_section().header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = f"{data['nom']} – {data['titre']}"
    header_paragraph.style = doc.styles['Heading 1']

    # Add contact information
    contact_table = doc.add_table(rows=1, cols=3)
    contact_table.style = 'Table Grid'
    contact_row = contact_table.rows[0].cells
    contact_row[0].text = f"📍 {data['ville']}" if data['ville'] else ""
    contact_row[1].text = f"📧 {data['email']}" if data['email'] else ""
    contact_row[2].text = f"📞 {data['telephone']}" if data['telephone'] else ""

    # Add sections
    def add_section(title, content):
        if content:
            doc.add_heading(title, level=2)
            doc.add_paragraph(content)

    add_section("Profil", data['profil'])
    add_section("Expériences professionnelles", data['experiences'])
    add_section("Compétences", data['competences'])
    add_section("Langues", data['langues'])
    add_section("Formation", data['formations'])
    add_section("Centres d'intérêt", data['interets'])

def generate_cv_fayssal_structure(doc, data):
    """Generate a CV using the 'cv_fayssal' structure."""
    # Add header with name, title, and photo
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    header_cells = header_table.rows[0].cells

    # Left cell: Name, title, and age
    header_cells[0].paragraphs[0].add_run(f"{data['nom']} {data['age']} ans").bold = True
    header_cells[0].paragraphs[0].add_run(f"\n{data['titre']}").italic = True

    # Right cell: Contact information
    contact_info = f"📧 {data['email']}\n📞 {data['telephone']}\n📍 {data['ville']}"
    header_cells[1].paragraphs[0].add_run(contact_info)

    # Add photo if available
    if 'photo_path' in data and data['photo_path']:
        try:
            header_cells[0].add_picture(data['photo_path'], width=docx.shared.Inches(1.5))
        except Exception as e:
            print(f"Erreur lors de l'ajout de la photo : {e}")

    # Add sections in two columns
    section_table = doc.add_table(rows=1, cols=2)
    section_table.style = 'Table Grid'
    left_column = section_table.rows[0].cells[0]
    right_column = section_table.rows[0].cells[1]

    # Left column: Profil, Compétences, Langues
    def add_section_to_column(column, title, content):
        if content:
            column.add_paragraph(title, style='Heading 2')
            column.add_paragraph(content)

    add_section_to_column(left_column, "Profil", data['profil'])
    add_section_to_column(left_column, "Compétences", data['competences'])
    add_section_to_column(left_column, "Langues", data['langues'])

    # Right column: Expériences professionnelles, Formation, Centres d'intérêt
    add_section_to_column(right_column, "Expériences professionnelles", data['experiences'])
    add_section_to_column(right_column, "Formation", data['formations'])
    add_section_to_column(right_column, "Centres d'intérêt", data['interets'])

@app.route('/download/<int:id>', methods=['GET'])
def download_cv(id):
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM cvs WHERE id = ?", (id,))
    row = cursor.fetchone()
    conn.close()

    # Generate the updated CV document
    doc = Document()
    generate_cv_fayssal_structure(doc, row)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"CV_{row['nom'].replace(' ', '_')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/edit_lm/<int:id>', methods=['GET', 'POST'])
def edit_lm(id):
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    if request.method == 'POST':
        nom = request.form.get("nom", "").strip()
        adresse = request.form.get("adresse", "").strip()
        ville = request.form.get("ville", "").strip()
        email = request.form.get("email", "").strip()
        tel = request.form.get("tel", "").strip()
        date_ville = request.form.get("date_ville", "").strip()
        sujet = request.form.get("sujet", "").strip()
        competences = request.form.get("competences", "").strip()
        experience = request.form.get("experience", "").strip()
        qualites = request.form.get("qualites", "").strip()
        cursor.execute('''
            UPDATE lettres
            SET nom=?, adresse=?, ville=?, email=?, tel=?, date_ville=?, sujet=?, competences=?, experience=?, qualites=?
            WHERE id=?
        ''', (nom, adresse, ville, email, tel, date_ville, sujet, competences, experience, qualites, id))
        conn.commit()
        conn.close()
        return "Lettre de motivation mise à jour avec succès.", 200
    cursor.execute("SELECT * FROM lettres WHERE id = ?", (id,))
    lettre = cursor.fetchone()
    conn.close()
    if not lettre:
        return "Lettre de motivation introuvable.", 404
    return render_template("edit_lm.html", lm=lettre)

@app.route('/download_lm/<int:id>', methods=['GET'])
def download_lm(id):
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM lettres WHERE id = ?", (id,))
    lm = cursor.fetchone()
    conn.close()
    if not lm:
        return "Lettre de motivation introuvable.", 404

    # Génération du contenu HTML pour la lettre
    lettre = f"""<b>{lm['nom']}</b><br><br><b>{lm['adresse']}</b><br><br><b>{lm['ville']}</b><br><br><b>{lm['email']}</b><br><br><b>Tél : {lm['tel']}</b><br><br><b>{lm['date_ville']}</b><br><br>
<b>Objet : Candidature pour un poste en {lm['sujet']}</b><br><br>
Madame, Monsieur,<br><br>
Je me permets de vous adresser ma candidature pour un poste dans le domaine du {lm['sujet']} au sein de votre entreprise.<br><br>
Titulaire de compétences telles que {lm['competences']}, """
    if lm['experience']:
        lettre += f"et fort(e) d'une expérience de {lm['experience']}, "
    lettre += f"""je souhaite mettre à profit mes connaissances et mon savoir-faire pour contribuer au développement de votre structure.<br><br>
Doté(e) de qualités personnelles comme {lm['qualites']}, je suis convaincu(e) de pouvoir m'intégrer rapidement à votre équipe et de participer activement à vos projets.<br><br>
Je reste à votre disposition pour un entretien afin de vous exposer plus en détail ma motivation et mon parcours.<br><br>
Je vous prie d’agréer, Madame, Monsieur, l’expression de mes salutations distinguées.<br><br>
{lm['nom']}"""

    htmlContent = f"""
    <html>
    <head><meta charset='utf-8'><title>Lettre de motivation</title></head>
    <body style="font-family:Arial,Helvetica,sans-serif;font-size:15px;color:#222;">
    {lettre}
    </body></html>
    """
    from io import BytesIO
    buffer = BytesIO()
    buffer.write(htmlContent.encode('utf-8'))
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Lettre_de_motivation_{lm['nom'].replace(' ', '_')}.doc",
        mimetype='application/msword'
    )

@app.route('/delete_lm/<int:id>', methods=['POST'])
def delete_lm(id):
    conn = sqlite3.connect("cv_data.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM lettres WHERE id = ?", (id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/delete_bulk', methods=['POST'])
def delete_bulk_cv():
    ids = request.form.getlist('ids')
    if ids:
        unique_ids = list(set(str(int(i)) for i in ids if i.isdigit()))
        if unique_ids:
            conn = sqlite3.connect("cv_data.db")
            cursor = conn.cursor()
            cursor.execute(
                f"DELETE FROM cvs WHERE id IN ({','.join(['?']*len(unique_ids))})",
                unique_ids
            )
            conn.commit()
            conn.close()
    return redirect(url_for('admin'))

@app.route('/delete_bulk_lm', methods=['POST'])
def delete_bulk_lm():
    ids = request.form.getlist('lm_ids')
    if ids:
        unique_ids = list(set(str(int(i)) for i in ids if i.isdigit()))
        if unique_ids:
            conn = sqlite3.connect("cv_data.db")
            cursor = conn.cursor()
            cursor.execute(
                f"DELETE FROM lettres WHERE id IN ({','.join(['?']*len(unique_ids))})",
                unique_ids
            )
            conn.commit()
            conn.close()
    return redirect(url_for('admin'))

@app.route('/delete_all_cvs', methods=['POST'])
def delete_all_cvs():
    conn = sqlite3.connect("cv_data.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM cvs")
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/delete_all_lettres', methods=['POST'])
def delete_all_lettres():
    conn = sqlite3.connect("cv_data.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM lettres")
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/download_modern/<int:id>', methods=['GET'])
def download_modern_cv(id):
    conn = sqlite3.connect("cv_data.db")
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM cvs WHERE id = ?", (id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return "CV introuvable.", 404

    # Rendu du template HTML moderne avec Jinja2
    rendered = render_template("cv_modern.html", cv=row)

    # Génération du PDF avec WeasyPrint
    pdf = HTML(string=rendered, base_url=request.base_url).write_pdf()

    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=CV_{row["nom"].replace(" ", "_")}_modern.pdf'
    return response

@app.route('/cv_modern_generator')
def cv_modern_generator():
    return render_template('cv_modern_generator.html')

@app.route('/cv_classic_editable')
def cv_classic_editable():
    return render_template('cv_classic_editable.html')

@app.route('/cv_creatif_generator')
def cv_creatif_generator():
    return render_template('cv_creatif_generator.html')

if __name__ == '__main__':
    # Conseils pour le lancement :
    # 1. Assurez-vous d'avoir installé toutes les dépendances :
    #    pip install flask python-docx pillow reportlab weasyprint
    # 2. Si WeasyPrint pose problème, installez les dépendances système :
    #    - Windows : https://weasyprint.readthedocs.io/en/stable/install.html#windows
    #    - Linux : sudo apt install libpango-1.0-0 libpangocairo-1.0-0 libcairo2 libgdk-pixbuf2.0-0 libffi-dev shared-mime-info
    # 3. Lancez l'app dans le terminal :
    #    python app.py
    # 4. Accédez à http://127.0.0.1:5000/ dans votre navigateur.

    app.run(debug=True)